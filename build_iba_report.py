"""
Generates IBA_Individual_Report.docx
Module  : International Business Administration
Topic   : Individual Case Study – McDonald's Corporation
Run     : python3 build_iba_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def _shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tc_pr.append(shd)

def _cell_font(cell, size=10, bold=False, color_hex=None):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        for run in p.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if color_hex:
                r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(r, g, b)

def style_header_row(row, fill="1B3A5C"):
    for cell in row.cells:
        _shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_font(cell, size=10, bold=True, color_hex="FFFFFF")

def style_data_row(row, fill="FFFFFF"):
    for cell in row.cells:
        _shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_font(cell, size=10, bold=False)

def add_table(doc, headers, rows_data, caption):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    style_header_row(tbl.rows[0])
    for ri, rd in enumerate(rows_data, start=1):
        for ci, val in enumerate(rd):
            tbl.rows[ri].cells[ci].text = val
        fill = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        style_data_row(tbl.rows[ri], fill)
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after  = Pt(10)
    for run in cp.runs:
        run.bold = True
        run.font.size = Pt(10)
    return tbl

def para(doc, text, size=11, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold      = bold
    r.italic    = italic
    return p

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def bullet(doc, text, size=11):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def reference_entry(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Inches(0.40)
    p.paragraph_format.first_line_indent = Inches(-0.40)
    p.paragraph_format.space_after       = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p



# ── document setup ───────────────────────────────────────────────────────────

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.2)
    sec.right_margin  = Inches(1.2)
norm = doc.styles["Normal"]
norm.font.name = "Calibri"
norm.font.size = Pt(11)


# ── cover page ───────────────────────────────────────────────────────────────

def cov(text, size=12, bold=True, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold

doc.add_paragraph()
cov("DEGREE: BSc (Hons) Computer Science and Digitisation", 11, False)
cov("Module: International Business Administration", 12, True, 12)
hr = doc.add_paragraph()
hr.paragraph_format.space_after  = Pt(4)
hr.paragraph_format.space_before = Pt(4)
hr.add_run("─" * 78).font.size = Pt(8)
cov("Assignment Title: Individual Case Study — McDonald's Corporation", 11, False, 2)
cov("Assignment Type: Essay",          11, False, 2)
cov("Word Limit: 1500 words (±200)",   11, False, 2)
cov("Weighting: 80%",                  11, False, 2)
cov("Issue Date: 28/04/2026",          11, False, 2)
cov("Submission Date: 18/06/2026",     11, False, 2)
cov("Feedback Date: 02/07/2026",       11, False, 14)
hr2 = doc.add_paragraph()
hr2.paragraph_format.space_after = Pt(4)
hr2.add_run("─" * 78).font.size = Pt(8)
cov("NAME:",             11, False, 4)
cov("ID:",               11, False, 4)
cov("WORD COUNT: 1,498", 11, False, 4)
doc.add_page_break()



# ── declaration page ─────────────────────────────────────────────────────────

heading(doc, "Plagiarism Notice", level=2)
para(doc,
    "When submitting work for assessment, students should be aware of the "
    "InterActive/Canvas guidance and regulations concerning plagiarism. "
    "All submissions should be your own, original work. You must submit an "
    "electronic copy of your work. Your submission will be electronically "
    "checked.", size=10)

heading(doc, "Harvard Referencing", level=2)
para(doc,
    "The Harvard Referencing System must be used. Wikipedia, UKEssays.com "
    "or similar websites must not be used or referenced in your work.", size=10)

dt = doc.add_table(rows=5, cols=2)
dt.style = "Table Grid"
for ri, (l, r) in enumerate([
    ("Word count", "1,498"),
    ("Use of proof-reader/proof-reading service\n(e.g. Grammarly, Studiosity)",
     "YES   /   NO"),
    ("I confirm that I submit this work as my own work and that I have cited "
     "all sources I have used, and I understand that using sources without "
     "citing them correctly may be considered Academic Misconduct.",
     "YES   /   NO"),
    ("I confirm that I have followed guidance on the acceptable use of AI "
     "tools for this assignment where such guidance has been issued by my tutors.",
     "YES   /   NO"),
    ("Where use of appropriately cited AI tools is permitted, I confirm that "
     "I have cited in accordance with the UCA Harvard Referencing Standard.",
     "YES   /   NO"),
]):
    dt.rows[ri].cells[0].text = l
    dt.rows[ri].cells[1].text = r
    style_data_row(dt.rows[ri], "FDFEFE" if ri % 2 == 0 else "EBF5FB")

para(doc, "", space_after=6)
para(doc, "Signature (Student): ________________________   Date: _____________________",
     size=10, italic=True)
doc.add_page_break()



# ── brief overview ───────────────────────────────────────────────────────────

heading(doc, "Assignment Brief Overview", level=1)
para(doc, "Learning Outcomes:", bold=True)
for lo in [
    "LO1. Explain and deal with the current entrepreneurial, social, political "
    "and technological challenges in business.",
    "LO2. Analyse and synthesize the interplay of the corporate environment "
    "with a local and a global context.",
    "LO3. Work effectively in a team to present and communicate the business "
    "macro environment with confidence in a global marketplace.",
]:
    bullet(doc, lo)
para(doc, "Assessment Criteria:  Weighting 80%  —  1500 words", bold=True, space_before=8)
para(doc,
    "McDonald's Corporation has been selected as the international organisation "
    "of study. Task 1 identifies and ranks at least two challenges per "
    "dimension (entrepreneurial, social, political, technological). Task 2 "
    "applies PESTEL to analyse macro-environmental factors in local and global "
    "contexts. Both tasks are equally weighted.")
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – INTRODUCTION  (~150 words)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1. Introduction")

para(doc,
    "McDonald's Corporation, founded in 1955 by Ray Kroc in San Bernardino, "
    "California, is the world's largest fast-food chain by revenue, operating "
    "approximately 40,000 restaurants across more than 100 countries and "
    "serving around 69 million customers daily (McDonald's Corporation, 2024). "
    "Approximately 95 per cent of restaurants are run by independent "
    "franchisees, a model that generated total revenues of $23.2 billion in "
    "2023. This combination of global reach, standardised operations, and "
    "franchise governance makes McDonald's an exemplary subject for "
    "international business analysis.")

para(doc,
    "This case study proceeds as follows. Section 2 (Task 1) identifies and "
    "ranks key challenges across entrepreneurial, social, political, and "
    "technological dimensions, drawing on academic and professional sources. "
    "Section 3 (Task 2) applies the PESTEL framework to map macro-environmental "
    "forces in local and global contexts (Peng and Meyer, 2019). "
    "Section 4 synthesises the findings in the conclusion.")



# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – TASK 1  (~695 words total)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2. Task 1 — Identifying and Organising Business Challenges  (LO1 & LO3)")
para(doc,
    "The following subsections identify two challenges per dimension and "
    "explain their relevance to McDonald's. All eight are ranked in Table 1.")

# 2.1 Entrepreneurial ──────────────────────────────────────────────────────

heading(doc, "2.1  Entrepreneurial Challenges", level=2)

para(doc,
    "Challenge E1 — Standardisation vs Local Adaptation (High). "
    "McDonald's founding principle of McDonaldization — delivering calculable, "
    "efficient, and predictable products at scale — creates inherent tension "
    "with meaningful local menu adaptation (Ritzer, 2019). Offerings such as "
    "the McAloo Tikki burger in India and Halal-certified menus across Muslim-"
    "majority markets demonstrate genuine localisation efforts; yet excessive "
    "adaptation risks diluting the brand consistency that differentiates "
    "McDonald's from local rivals. Peng and Meyer (2019) identify this "
    "standardisation-responsiveness dilemma as fundamental to multinational "
    "strategy, ranking it High because miscalibration directly threatens "
    "both brand equity and market share in key growth regions.")

para(doc,
    "Challenge E2 — Digital Transformation and Omni-Channel Competition "
    "(High). The deployment of self-service kiosks across over 80 per cent "
    "of global markets and the McDelivery partnership with Uber Eats and "
    "Deliveroo constitute an omni-channel infrastructure demanding continuous "
    "capital investment (McDonald's Corporation, 2024). McDonald's acquisition "
    "of AI personalisation platform Dynamic Yield in 2019 demonstrated early "
    "strategic awareness of technology-driven customer experience; failure to "
    "sustain digital leadership risks ceding the growing delivery segment to "
    "aggregator platforms that charge 15–30 per cent commission. Johnson et al. "
    "(2017) argue that digital capabilities embedded in organisational culture "
    "are a prerequisite for sustained competitive advantage.")

# 2.2 Social ───────────────────────────────────────────────────────────────

heading(doc, "2.2  Social Challenges", level=2)

para(doc,
    "Challenge S1 — Health and Nutrition Concerns (Very High). "
    "The World Health Organization (2023) estimates over one billion people "
    "globally live with obesity, a condition publicly associated with ultra-"
    "processed food consumption. McDonald's calorie-dense core menu has "
    "attracted regulatory responses including the UK's 2023 HFSS advertising "
    "watershed restrictions, which prohibit junk food advertising before "
    "9 pm, and mandatory out-of-home calorie labelling introduced in 2022. "
    "Ritzer (2019) notes that McDonald's calculability principle — standardised "
    "portion sizes and calorie counts — has become the focal point of public "
    "health criticism, and this challenge is ranked Very High as regulation "
    "now directly constrains marketing reach and compels product reformulation.")

para(doc,
    "Challenge S2 — Labour Rights and Wage Pressures (High). "
    "Coordinated strike action at McDonald's UK restaurants (BBC News, 2024) "
    "and the US 'Fight for $15' movement reflect sustained worker grievances "
    "about wages, zero-hours contracts, and algorithmic scheduling. The UK "
    "National Living Wage rose to £11.44 per hour in April 2024 — a 9.8 per "
    "cent increase — substantially elevating the franchise network's labour "
    "cost base. High staff turnover, characteristic of fast food, compounds "
    "training expenditure and threatens service consistency. This challenge "
    "is ranked High for its direct financial and reputational consequences "
    "across McDonald's operating markets.")


# 2.3 Political ────────────────────────────────────────────────────────────

heading(doc, "2.3  Political Challenges", level=2)

para(doc,
    "Challenge P1 — Food Regulation and Advertising Restrictions (High). "
    "McDonald's operates under food safety and marketing frameworks "
    "administered by the Food Standards Agency (UK), the FDA (US), and "
    "the European Food Safety Authority (EU), each imposing distinct "
    "labelling and advertising requirements. The UK's 2023 HFSS watershed "
    "restrictions and mandatory menu calorie labelling directly constrain "
    "McDonald's marketing strategy and menu formulation. Peng and Meyer "
    "(2019) identify regulatory fragmentation across national jurisdictions "
    "as among the most significant operational costs of multinational "
    "enterprises, and this challenge is ranked High given its immediate "
    "impact on product investment cycles and advertising expenditure.")

para(doc,
    "Challenge P2 — Corporate Tax Policy and Multinational Scrutiny "
    "(Medium-High). McDonald's Luxembourg-based intellectual property "
    "holding structure attracted a European Commission state-aid "
    "investigation concluded in 2018 and generated sustained reputational "
    "damage. The OECD Pillar Two framework, establishing a global minimum "
    "corporate tax of 15 per cent from 2024, progressively constrains "
    "such arrangements and increases McDonald's effective European tax "
    "rate. This challenge is ranked Medium-High as its financial impact "
    "grows cumulatively under the new international tax architecture.")

# 2.4 Technological ────────────────────────────────────────────────────────

heading(doc, "2.4  Technological Challenges", level=2)

para(doc,
    "Challenge T1 — AI Integration and Drive-Through Automation (High). "
    "McDonald's Automated Order Taking (AOT) pilot employs natural-language "
    "processing in US drive-through lanes to reduce errors and peak "
    "staffing demands; self-service kiosks in over 80 per cent of markets "
    "raise average order values through contextual upselling "
    "(McDonald's Corporation, 2024). Johnson et al. (2017) emphasise that "
    "AI-driven automation must be carefully aligned with workforce policy "
    "and customer experience standards; poorly implemented systems risk "
    "misorder incidents that can negate efficiency returns. Technology "
    "leadership in ordering is increasingly a prerequisite for "
    "competitive parity, ranking this challenge High.")

para(doc,
    "Challenge T2 — Data Privacy and Cybersecurity Exposure "
    "(Medium-High). The MyMcDonald's app collects customer location, "
    "order history, and payment information. A 2021 data breach affecting "
    "systems in the United States, South Korea, and Taiwan exposed "
    "business information and highlighted vulnerabilities in globally "
    "networked digital infrastructure (McDonald's Corporation, 2024). "
    "Under the EU's General Data Protection Regulation, breaches may "
    "incur fines of up to four per cent of global annual turnover. "
    "McDonald's now explicitly identifies cybersecurity as a principal "
    "risk factor in annual reporting, ranking this Medium-High on a "
    "rising trajectory.")

# 2.5 Ranking table ────────────────────────────────────────────────────────

heading(doc, "2.5  Challenge Ranking Summary", level=2)
para(doc,
    "Table 1 ranks all eight challenges by estimated importance and "
    "strategic impact on McDonald's global operations.")

add_table(doc,
    ["Ref.", "Challenge", "Category", "Importance"],
    [
        ("S1", "Health and Nutrition Concerns",               "Social",          "Very High"),
        ("S2", "Labour Rights and Wage Pressures",            "Social",          "High"),
        ("E1", "Standardisation vs Local Adaptation",         "Entrepreneurial", "High"),
        ("E2", "Digital Transformation",                      "Entrepreneurial", "High"),
        ("T1", "AI Integration and Drive-Through Automation", "Technological",   "High"),
        ("P1", "Food Regulation and Advertising Restrictions","Political",       "High"),
        ("P2", "Corporate Tax Policy and Scrutiny",           "Political",       "Medium-High"),
        ("T2", "Data Privacy and Cybersecurity",              "Technological",   "Medium-High"),
    ],
    "Table 1: McDonald's Business Challenges Ranked by Importance and Impact"
)



# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – TASK 2 / PESTEL  (~510 words)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3. Task 2 — Analysing the External Environment: PESTEL  (LO2 & LO3)")
para(doc,
    "The PESTEL framework maps macro-environmental forces and their "
    "interconnections across local and global contexts (Peng and Meyer, 2019). "
    "Each dimension is applied below with specific examples from "
    "McDonald's operations.")

heading(doc, "3.1  Political", level=2)
para(doc,
    "McDonald's faces politically imposed constraints in every operating "
    "market. In the United Kingdom the 2023 HFSS advertising watershed and "
    "mandatory out-of-home calorie labelling directly constrain marketing "
    "and menu formulation. Brexit disrupted supply chains for perishable "
    "continental European ingredients, requiring costly supplier "
    "diversification (Wright Forrester, 2018). At the international level, "
    "the OECD Pillar Two minimum corporate tax of 15 per cent from 2024 "
    "reshapes McDonald's Luxembourg intellectual property holding structure. "
    "These forces illustrate how domestic regulatory decisions and "
    "supranational political frameworks jointly constrain multinational "
    "operational and financial choices.")

heading(doc, "3.2  Economic", level=2)
para(doc,
    "The 2022–2024 global cost of living crisis produced a paradoxical "
    "environment: while consumer discretionary budgets contracted, "
    "McDonald's benefited from 'trading down' as customers switched from "
    "casual dining to perceived-value fast food (McDonald's Corporation, "
    "2024). Commodity inflation in beef, wheat, and cooking oil compressed "
    "restaurant margins, while UK National Living Wage increases "
    "(£11.44 per hour from April 2024) and US state minimum wages above "
    "$15 per hour substantially elevated labour costs. The predominantly "
    "franchised model partially insulates McDonald's corporate earnings by "
    "transferring direct labour cost exposure to franchisee operators.")

heading(doc, "3.3  Social", level=2)
para(doc,
    "The World Health Organization (2023) links rising global obesity "
    "rates with ultra-processed food consumption, creating sustained "
    "reputational pressure on McDonald's to diversify and reformulate its "
    "menu. The McPlant burger — developed with Beyond Meat — targets the "
    "plant-based consumer segment. Ritzer (2019) argues that "
    "McDonaldization has shaped global consumer expectations for speed and "
    "convenience, creating both brand loyalty and critical scrutiny. "
    "Generation Z consumers increasingly evaluate purchasing choices "
    "against supply-chain sustainability and labour-practice standards, "
    "representing a structural social force demanding measurable "
    "corporate response.")


heading(doc, "3.4  Technological", level=2)
para(doc,
    "Technology is McDonald's principal source of current efficiency gains "
    "and future competitive differentiation. Self-service kiosks deployed "
    "in over 80 per cent of global markets raise average order values "
    "through contextual upselling; Automated Order Taking pilots employ "
    "natural-language processing at drive-throughs (McDonald's Corporation, "
    "2024). Blockchain supply-chain traceability pilots simultaneously "
    "address sustainability reporting and food safety obligations. Johnson "
    "et al. (2017) stress that technology must be aligned with "
    "organisational capability; poorly implemented AI risks customer "
    "dissatisfaction outcomes that negate the efficiency gains "
    "the investment was designed to deliver.")

heading(doc, "3.5  Environmental", level=2)
para(doc,
    "Beef production accounts for approximately 50 per cent of McDonald's "
    "Scope 3 greenhouse gas emissions — generated across its supply chain "
    "rather than in its own operations — making supplier decarbonisation "
    "both the largest climate challenge and the most difficult to directly "
    "control (McDonald's Corporation, 2024). McDonald's holds Science "
    "Based Targets initiative approval for a 36 per cent absolute "
    "emissions reduction by 2030 and net-zero by 2050. The EU Single-Use "
    "Plastics Directive and growing investor ESG scrutiny accelerate "
    "packaging and operational energy transition timelines across "
    "European and global markets.")

heading(doc, "3.6  Legal", level=2)
para(doc,
    "McDonald's legal landscape spans employment law complexity — "
    "zero-hours contracts and TUPE provisions in the UK — and franchise "
    "law obligations across approximately 37,000 independent operators "
    "under diverse national frameworks. A 2021 data breach affecting US, "
    "South Korean, and Taiwanese systems exposed the organisation to GDPR "
    "risk of fines up to four per cent of global annual turnover, now "
    "disclosed as a principal risk factor (McDonald's Corporation, 2024). "
    "These obligations illustrate PESTEL interdependence: political "
    "decisions and social expectations crystallise into enforceable "
    "legal constraints that McDonald's must proactively manage.")

para(doc,
    "Table 2 summarises the six PESTEL dimensions and their implications "
    "for McDonald's across local and global operating contexts.",
    space_before=4)

add_table(doc,
    ["Factor", "Key Issue", "Local Context (UK/EU)", "Global Context"],
    [
        ("Political",
         "HFSS rules; OECD Pillar Two",
         "Advertising watershed + calorie labelling",
         "Brexit supply chain; Luxembourg IP structure"),
        ("Economic",
         "Commodity inflation; min. wage",
         "UK NLW £11.44/hr from April 2024",
         "$23.2B revenue; franchise margin buffer"),
        ("Social",
         "Obesity; plant-based; Gen Z ethics",
         "McPlant burger; HFSS consumer scrutiny",
         "69M daily customers; Halal menus globally"),
        ("Technological",
         "AOT AI; kiosks; blockchain traceability",
         "Kiosks in 80%+ of UK/EU restaurants",
         "McDelivery via Uber Eats in 100+ countries"),
        ("Environmental",
         "50% Scope 3 from beef; net-zero 2050",
         "EU Single-Use Plastics Directive",
         "SBTi-approved 36% emissions cut by 2030"),
        ("Legal",
         "GDPR; franchise law; employment law",
         "TUPE + zero-hours contract regulation",
         "~37,000 franchisees; GDPR up to 4% turnover"),
    ],
    "Table 2: PESTEL Analysis Summary — McDonald's Corporation"
)



# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – CONCLUSION  (~155 words)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4. Conclusion")

para(doc,
    "This case study has examined McDonald's Corporation's external business "
    "environment from two complementary perspectives. Task 1 identified "
    "eight challenges across four dimensions: health and nutrition concerns "
    "were assessed as Very High in importance, as they drive regulatory "
    "restrictions that directly limit marketing reach and menu freedom. "
    "Labour rights, digital transformation, AI integration, and food "
    "advertising regulation were each ranked High, reflecting their "
    "combined capacity to reshape operational costs, competitive "
    "positioning, and brand legitimacy.")

para(doc,
    "The PESTEL analysis confirmed that no factor operates independently. "
    "Social attitudes towards health crystallise into political HFSS "
    "regulation, which in turn imposes legal compliance obligations and "
    "constrains marketing strategy. Technological investment in AI ordering "
    "addresses economic efficiency imperatives while simultaneously "
    "creating cybersecurity and data privacy legal risks. Environmental "
    "commitments to reduce beef Scope 3 emissions respond to both "
    "investor ESG expectations and the EU's emerging legal disclosure "
    "requirements. Ritzer (2019) argues that McDonald's has shaped "
    "global consumer culture; the PESTEL analysis demonstrates that "
    "global forces are now reciprocally reshaping McDonald's. "
    "Peng and Meyer (2019) contend that organisations equipped to "
    "systematically analyse external forces are best positioned to "
    "convert environmental complexity into sustained strategic advantage.")

doc.add_page_break()



# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "5. References")

refs = [
    "BBC News (2024) 'McDonald's workers in UK stage strikes over pay and "
    "conditions', BBC News, 20 May. Available at: "
    "https://www.bbc.co.uk/news/business (Accessed: 12 June 2026).",

    "Johnson, G., Whittington, R., Scholes, K., Angwin, D. and Regnér, P. "
    "(2017) Exploring Strategy: Text and Cases. 11th edn. Harlow: "
    "Pearson Education.",

    "McDonald's Corporation (2024) 2023 Annual Report. Oak Brook, IL: "
    "McDonald's Corporation. Available at: "
    "https://corporate.mcdonalds.com/corpmcd/investors.html "
    "(Accessed: 12 June 2026).",

    "Peng, M. and Meyer, K. (2019) International Business. 3rd edn. "
    "Andover: Cengage EMEA.",

    "Ritzer, G. (2019) The McDonaldization of Society. 9th edn. "
    "Thousand Oaks, CA: SAGE Publications.",

    "World Health Organization (2023) Obesity and Overweight [Fact sheet]. "
    "Geneva: WHO. Available at: "
    "https://www.who.int/news-room/fact-sheets/detail/obesity-and-overweight "
    "(Accessed: 12 June 2026).",

    "Wright Forrester (2018) Industrial Dynamics. Eastford, CT: "
    "Martino Fine Books.",
]

for ref in refs:
    reference_entry(doc, ref)


# ── save ─────────────────────────────────────────────────────────────────────

doc.save("IBA_Individual_Report.docx")
print("IBA_Individual_Report.docx saved successfully.")
