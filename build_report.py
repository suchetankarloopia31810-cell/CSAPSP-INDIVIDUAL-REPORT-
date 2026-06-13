"""
Generates CSAPSP_Individual_Report.docx
Topic : Hospital Emergency Department Patient Triage Management System
Run   : python3 build_report.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ════════════════════════════════════════════════════════════════════════════
#  STYLE HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tc_pr.append(shd)


def _cell_font(cell, size=10, bold=False, color_hex=None):
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        for run in para.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if color_hex:
                r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(r, g, b)


def style_header_row(row, fill="1B3A5C"):
    for cell in row.cells:
        _shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_font(cell, size=10, bold=True, color_hex="FFFFFF")


def style_data_row(row, fill="FFFFFF"):
    for cell in row.cells:
        _shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_font(cell, size=10, bold=False)


def add_table(doc, headers, rows_data, caption):
    """Build a styled table with dark header and alternating rows."""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    style_header_row(tbl.rows[0])

    # data rows
    for ri, row_data in enumerate(rows_data, start=1):
        for ci, val in enumerate(row_data):
            tbl.rows[ri].cells[ci].text = val
        fill = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        style_data_row(tbl.rows[ri], fill)

    # caption
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after  = Pt(10)
    for run in cp.runs:
        run.bold = True
        run.font.size = Pt(10)
    return tbl


def para(doc, text, size=11, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=0, space_after=6, left_indent=None, first_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent  is not None: p.paragraph_format.left_indent        = left_indent
    if first_indent is not None: p.paragraph_format.first_line_indent  = first_indent
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F4F6F7")
    pPr.append(shd)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def figure_caption(doc, text):
    cp = doc.add_paragraph(text)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after  = Pt(12)
    for run in cp.runs:
        run.bold = True
        run.font.size = Pt(10)
    return cp


def reference_entry(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Inches(0.40)
    p.paragraph_format.first_line_indent = Inches(-0.40)
    p.paragraph_format.space_after       = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


# ════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── page setup ───────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.2)
    sec.right_margin  = Inches(1.2)

# ── default style ────────────────────────────────────────────────────────────
norm = doc.styles["Normal"]
norm.font.name = "Calibri"
norm.font.size = Pt(11)


# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

def cover_line(text, size=12, bold=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p

doc.add_paragraph()
cover_line("DEGREE: Computer Science and Digitisation", 11, False)
cover_line("Module: Algorithms and Problem Solving using Python", 12, True, 12)

# horizontal rule
hr = doc.add_paragraph()
hr.paragraph_format.space_after  = Pt(4)
hr.paragraph_format.space_before = Pt(4)
hr_run = hr.add_run("─" * 78)
hr_run.font.size = Pt(8)

cover_line("Assignment Title:", 11, False, 2)
cover_line(
    "Prioritising Lives: A Python-Based Sorting Algorithm Analysis\n"
    "for Hospital Emergency Department Patient Triage Management",
    13, True, 14)

cover_line("Assignment Type: Individual Report", 11, False, 2)
cover_line("Word Limit: 2000 words (±200)",     11, False, 2)
cover_line("Weighting: 50%",                    11, False, 2)
cover_line("Issue Date: 28/04/2026",            11, False, 2)
cover_line("Submission Date: 15/06/2026",       11, False, 2)
cover_line("Feedback Date: 29/06/2026",         11, False, 2)
cover_line("Issued by: Dr. Syed Arslan Abbas Rizvi", 11, False, 14)

hr2 = doc.add_paragraph()
hr2.paragraph_format.space_after  = Pt(4)
hr2_run = hr2.add_run("─" * 78)
hr2_run.font.size = Pt(8)

cover_line("NAME:",       11, False, 4)
cover_line("ID:",         11, False, 4)
cover_line("WORD COUNT:  2,053", 11, False, 4)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  PLAGIARISM & LEARNER DECLARATION  (matching BSBI template page 2)
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "Plagiarism Notice", level=2)
para(doc,
    "When submitting work for assessment, students should be aware of the "
    "InterActive/Canvas guidance and regulations concerning plagiarism. All "
    "submissions should be your own, original work. You must submit an "
    "electronic copy of your work. Your submission will be electronically "
    "checked.", size=10)

heading(doc, "Harvard Referencing", level=2)
para(doc,
    "The Harvard Referencing System must be used. Wikipedia, UKEssays.com or "
    "similar websites must not be used or referenced in your work.", size=10)

# Learner declaration table
decl_tbl = doc.add_table(rows=4, cols=2)
decl_tbl.style = "Table Grid"
decl_data = [
    ("Word count", ""),
    ("Use of proof-reader/proof-reading service\n(e.g. Grammarly, Studiosity)", "YES   /   NO"),
    ("I confirm that I submit this work as my own work and that I have cited "
     "all sources I have used, and I understand that using sources without "
     "citing them correctly may be considered Academic Misconduct.", "YES   /   NO"),
    ("I confirm that I have followed guidance on the acceptable use of AI "
     "tools for this assignment where such guidance has been issued by my "
     "tutors.", "YES   /   NO"),
]
for ri, (left, right) in enumerate(decl_data):
    decl_tbl.rows[ri].cells[0].text = left
    decl_tbl.rows[ri].cells[1].text = right
    style_data_row(decl_tbl.rows[ri], "FDFEFE" if ri % 2 == 0 else "EBF5FB")

para(doc, "", space_after=6)
para(doc, "Signature (Student): ________________________   Date: _____________________",
     size=10, italic=True)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION PAGE (Learning Outcomes + Assessment Criteria)
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "Introduction", level=1)
para(doc, "Learning Outcomes:", bold=True)

los = [
    "LO1. Demonstrate an understanding and examine how different data structures "
    "and algorithm design methods including lists, stacks, queues, trees, and graphs "
    "impact the performance of programs.",
    "LO2. Implement the algorithms for problem solving using Python code.",
    "LO3. Implement and execute Python programs using functions, modules, libraries "
    "and classes.",
]
for lo in los:
    p = doc.add_paragraph(lo, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)

para(doc, "Assessment Criteria:  Weighting 50%  —  2000 words", bold=True, space_before=8)
para(doc,
    "The objective of this assignment is to demonstrate proficiency in algorithms, "
    "data structures, and problem-solving techniques using Python. Through practical "
    "implementations, students explore various data structures, algorithm design "
    "methods, and Python programming features to solve real-world problems effectively.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "1. Introduction")

para(doc,
    "Hospital emergency departments (EDs) are among the most time-critical "
    "environments in modern healthcare. Every year, NHS emergency departments in "
    "England handle tens of millions of patient attendances, with clinical "
    "outcomes that are directly influenced by how quickly the most seriously ill "
    "patients receive attention (NHS England, 2024). A fundamental tool that "
    "enables this prioritisation is the Manchester Triage System (MTS), a "
    "five-level scoring framework in which trained nurses assign each arriving "
    "patient a score from 1 (Immediate — life-threatening) to 5 (Non-Urgent — "
    "can wait) based on presenting symptoms (Mackway-Jones, Marsden and Windle, "
    "2014). Managing and re-ordering this queue efficiently, as patient volumes "
    "fluctuate throughout a shift, is a non-trivial computational problem.")

para(doc,
    "Algorithms sit at the heart of such problems. The choice of sorting "
    "algorithm determines whether a system responds in milliseconds or minutes "
    "when the patient census climbs from a quiet overnight shift to a peak "
    "Saturday afternoon surge (Cormen et al., 2022). A poorly chosen O(n²) "
    "algorithm that processes fifty records in under a millisecond may take "
    "more than eleven seconds on ten thousand records — an unacceptable delay "
    "in a safety-critical setting.")

para(doc,
    "This project transforms the ED triage problem into a concrete algorithmic "
    "challenge implemented in Python. Four classical sorting algorithms are "
    "designed, coded and compared: Bubble Sort, Insertion Sort, Merge Sort, "
    "and a three-way partitioned Quick Sort. A custom OpCounter class measures "
    "every comparison, swap and arithmetic operation performed, allowing "
    "empirical operation counts to be validated against theoretical Big-O "
    "predictions. A flag_critical() function scans the sorted queue and "
    "identifies all patients at severity Level 1 or Level 2, who require "
    "immediate clinical intervention (Das, 2025).")

para(doc,
    "The system is tested across three dataset scales — a seven-patient shift "
    "handover log, a fifty-patient mid-shift snapshot, and a ten-thousand-record "
    "monthly audit export — providing a rigorous basis for algorithm selection "
    "across the full operational range of an NHS emergency department.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – PROBLEM ANALYSIS
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "2. Problem Analysis and Requirements")

para(doc,
    "The scenario centres on an NHS emergency department that logs patient "
    "arrivals in a plain-text file. Each record contains two fields: the "
    "patient's name (a string) and their MTS triage score (an integer from "
    "1 to 5). A representative sample of seven records from the assignment "
    "scenario is shown in Table 0 below.")

add_table(doc,
    ["Patient Name", "Triage Score", "MTS Category"],
    [
        ("James",  "3", "Urgent (Yellow)"),
        ("Maria",  "1", "Immediate (Red)"),
        ("Chen",   "4", "Standard (Green)"),
        ("Fatima", "2", "Very Urgent (Orange)"),
        ("Oliver", "5", "Non-Urgent (Blue)"),
        ("Priya",  "1", "Immediate (Red)"),
        ("Marcus", "3", "Urgent (Yellow)"),
    ],
    "Table 0: Assignment Sample Patient Dataset with MTS Categories"
)

para(doc,
    "The primary functional requirement is to sort these records in ascending "
    "order of triage score, so that the patient with the lowest (most critical) "
    "score is placed first in the output queue. Secondary requirements include: "
    "flagging all patients with a score of 1 or 2 as critical; reporting the "
    "total number of records processed; measuring sorting time in milliseconds; "
    "and counting every comparison, swap and arithmetic operation performed "
    "(Mackway-Jones, Marsden and Windle, 2014).")

para(doc,
    "A non-functional requirement is scalability. The system must process both "
    "small shift logs of around seven to twenty patients and large monthly "
    "exports of up to ten thousand records without unacceptable degradation in "
    "response time. This constraint makes algorithmic complexity a primary "
    "selection criterion. The input data is read from a text file or generated "
    "synthetically using Python's random module; either way, each record is "
    "stored as a Python dictionary before sorting begins (Das, 2025).")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – DATA STRUCTURES
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "3. Data Structures and Data Types Selection")

para(doc,
    "Selecting an appropriate data structure is as consequential as choosing the "
    "right algorithm; together they govern memory usage, access patterns and "
    "overall throughput (Fatima, 2023). This system uses Python's built-in types "
    "arranged in a composite structure to represent the patient queue.")

para(doc,
    "Each patient's name is stored as a Python string, which supports "
    "the character-level access and display operations needed for output. The "
    "triage score is stored as a Python integer, enabling the numeric comparisons "
    "that drive the sort key. Both fields are encapsulated together in a "
    "dictionary using the keys 'name' and 'score', for example:")

code_block(doc,
    "patients = [\n"
    "    {'name': 'Maria',  'score': 1},\n"
    "    {'name': 'James',  'score': 3},\n"
    "    {'name': 'Oliver', 'score': 5},\n"
    "]")

para(doc,
    "The outer container is a Python list, which provides zero-based integer "
    "indexing, in-place element swapping, and dynamic resizing. These properties "
    "are precisely what comparison-based sorting algorithms require: the ability "
    "to read any element in O(1) time and exchange two elements without "
    "reallocating memory (Lafore, Broder and Canning, 2022). A list of "
    "dictionaries is therefore the most natural fit for this problem.")

para(doc,
    "Alternative structures were considered but rejected for this task. A stack "
    "follows Last-In-First-Out (LIFO) semantics, which is useful for undo "
    "operations or depth-first graph traversal; it provides no mechanism to "
    "maintain sorted order. A queue's First-In-First-Out (FIFO) structure suits "
    "patient arrival logging but cannot re-order records by triage score. A "
    "Binary Search Tree (BST) maintains sorted order and supports O(log n) "
    "queries, but constructing the tree introduces O(n log n) overhead that "
    "offers no advantage over simply sorting a list, and it adds significant "
    "implementation complexity (Fatima, 2023). A graph models relational "
    "networks and has no natural application to a linear sequence of patient "
    "records. The list-of-dictionaries structure is therefore the optimal choice "
    "for this clinical triage scenario.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – ALGORITHM ANALYSIS
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "4. Sorting Algorithm Analysis")

para(doc,
    "Four sorting algorithms were evaluated against the triage system's "
    "requirements. Table 1 presents their theoretical Big-O time and space "
    "complexity across best, average and worst-case scenarios.")

add_table(doc,
    ["Algorithm", "Best Case", "Average Case", "Worst Case", "Space"],
    [
        ("Bubble Sort",              "O(n)",       "O(n²)",       "O(n²)",       "O(1)"),
        ("Insertion Sort",           "O(n)",       "O(n²)",       "O(n²)",       "O(1)"),
        ("Merge Sort",               "O(n log n)", "O(n log n)",  "O(n log n)",  "O(n)"),
        ("Quick Sort (3-way pivot)", "O(n)",       "O(n log n)",  "O(n log n)*", "O(log n)"),
    ],
    "Table 1: Big-O Complexity Comparison  (*3-way partition avoids O(n²) on repeated values)"
)

para(doc,
    "Bubble Sort scans the list repeatedly, swapping adjacent pairs that are "
    "out of triage order. Each complete pass moves the least-urgent patient to "
    "the end of the unsorted section. While straightforward to implement, its "
    "O(n²) average complexity means it requires approximately 49.9 million "
    "comparisons on ten thousand records — wholly impractical for a live "
    "clinical environment (Sabah et al., 2023).")

para(doc,
    "Insertion Sort processes the queue one patient at a time, inserting each "
    "new arrival into the correct position among those already sorted. It "
    "achieves O(n) on nearly-sorted input, which makes it attractive for "
    "small shift-handover logs where patients arrive in roughly chronological "
    "order. However, its O(n²) worst case renders it unsuitable for large "
    "monthly audit exports (Shabbir et al., 2023).")

para(doc,
    "Merge Sort uses a divide-and-conquer strategy: it recursively halves the "
    "patient list until each sub-list contains one element, then merges pairs "
    "of sub-lists by comparing triage scores. Its time complexity is O(n log n) "
    "in all cases — best, average and worst — making its performance entirely "
    "predictable regardless of the order in which patients arrive. The "
    "trade-off is an O(n) auxiliary space requirement for the temporary arrays "
    "created during the merge step (Cormen et al., 2022).")

para(doc,
    "Standard Quick Sort (Lomuto partition) degenerates to O(n²) when many "
    "records share the same key value. With triage scores limited to just five "
    "distinct integers, this is not a theoretical edge case but a guaranteed "
    "outcome: testing revealed a RecursionError at 10,000 records due to "
    "stack depth exceeding 995 recursive calls. This limitation was resolved "
    "by implementing a three-way Dutch National Flag partition with randomised "
    "pivot selection, which groups equal scores together in a single pass and "
    "eliminates the degenerate recursion pattern (Cormen et al., 2022). The "
    "result is an expected O(n log n) complexity that, with only five distinct "
    "values, achieves O(n) best-case performance on already-grouped data.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – FLOWCHART
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "5. Algorithm Design and Flowchart Explanation")

para(doc,
    "The complete algorithm is described by the flowchart in Figure 1. Each "
    "shape follows standard flowchart conventions: ovals denote terminal nodes "
    "(START and END), rectangles represent process steps, and diamonds indicate "
    "decision points with Yes/No branches. The flowchart covers all stages from "
    "initial data loading through to final output, including the critical-patient "
    "flag path.")

# ── Embed flowchart image ────────────────────────────────────────────────────
if os.path.exists("flowchart.png"):
    doc.add_picture("flowchart.png", width=Inches(5.8))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
figure_caption(doc, "Figure 1: Hospital ED Patient Triage System — Algorithm Flowchart")

para(doc,
    "The algorithm begins by loading patient records from the input file or "
    "generating a synthetic test dataset. A validation step checks that every "
    "record contains a string name and an integer triage score in the range "
    "1–5; any malformed entry is skipped and logged. Valid records are stored "
    "in the list-of-dictionaries structure and the OpCounter is reset to zero "
    "before the high-resolution timer starts (Reya et al., 2023).")

para(doc,
    "The Merge Sort stage follows the left branch of the 'len(data) > 1?' "
    "decision. The list is split at its midpoint into a left half and a right "
    "half; each half is sorted recursively until the base case of a "
    "single-element list is reached. The merge step then reconstructs the "
    "sorted list by repeatedly comparing the leading element of each half and "
    "appending the smaller score to the output, incrementing comparison and "
    "arithmetic counters throughout. Once sorting completes, the timer stops "
    "and elapsed time is computed. The sorted queue is displayed, and the "
    "decision diamond 'score ≤ 2?' filters every Level-1 and Level-2 patient "
    "into a separate critical-flag output before the programme prints the "
    "final performance summary and terminates.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – PYTHON IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "6. Python Implementation")

para(doc,
    "Python was selected for this implementation because of its expressive "
    "syntax, comprehensive standard library and native support for the list "
    "and dictionary types central to this design. The programme is structured "
    "into four components: standard library imports, the OpCounter class, the "
    "sorting functions, and the main execution block.")

para(doc,
    "Three standard library modules are used. The time module provides "
    "time.perf_counter(), a high-resolution clock that measures elapsed time "
    "with sub-microsecond precision and is the recommended approach for "
    "benchmarking short code segments in Python (Reya et al., 2023). The "
    "random module generates synthetic patient datasets with uniformly "
    "distributed triage scores in the range 1–5. The copy module's deepcopy() "
    "function ensures that each sorting algorithm receives an independent copy "
    "of the original list, so no algorithm's in-place mutations affect "
    "subsequent tests.")

para(doc,
    "The OpCounter class separates measurement logic cleanly from sorting "
    "logic. Each sorting function receives an OpCounter instance and increments "
    "its comparisons, swaps and arithmetic attributes at every corresponding "
    "operation. This approach provides granular visibility into algorithmic "
    "work beyond what execution time alone reveals (Abuba et al., 2025):")

code_block(doc,
    "class OpCounter:\n"
    "    def __init__(self):\n"
    "        self.comparisons = 0\n"
    "        self.swaps       = 0\n"
    "        self.arithmetic  = 0\n"
    "\n"
    "    def total(self):\n"
    "        return self.comparisons + self.swaps + self.arithmetic")

para(doc,
    "The Merge Sort implementation consists of a public merge_sort() wrapper "
    "and a recursive _merge_helper() function. The helper splits the input "
    "list at the midpoint, recursively sorts each half, then reconstructs the "
    "sorted list by comparing 'score' fields from the two halves and copying "
    "the smaller value into the output position. The core merge loop is shown "
    "below:")

code_block(doc,
    "while i < len(left) and j < len(right):\n"
    "    counter.comparisons += 1\n"
    "    if left[i]['score'] <= right[j]['score']:\n"
    "        data[k] = left[i];  i += 1\n"
    "    else:\n"
    "        data[k] = right[j]; j += 1\n"
    "    k += 1\n"
    "    counter.arithmetic += 1")

para(doc,
    "The flag_critical() function accepts the sorted patient list and a "
    "threshold (defaulting to 2) and returns all patients whose triage score "
    "is at or below that threshold using a Python list comprehension. Because "
    "the input is already sorted in ascending order, all critical patients "
    "appear as a contiguous block at the start of the list, making the output "
    "both correct and easy for clinical staff to act upon (Das, 2025). "
    "Figure 2 shows the programme's console output for the small dataset.")

code_block(doc,
    "HOSPITAL ED PATIENT TRIAGE MANAGEMENT SYSTEM\n"
    "============================================================\n"
    "[SMALL DATASET]  7 patients\n"
    "Algorithm            Time (ms)  Comparisons  Swaps  Arithmetic  Total Ops\n"
    "---------------------------------------------------------------------------\n"
    "Bubble Sort             0.0460           21      9          28         58\n"
    "Insertion Sort          0.0333           14      9          20         43\n"
    "Merge Sort              0.0406           14      0          40         54\n"
    "Quick Sort (3-way)      0.0423           15     13          37         65\n"
    "\n"
    "Sorted queue (ascending triage score):\n"
    "  Maria    Score 1 – Immediate   (Red)\n"
    "  Priya    Score 1 – Immediate   (Red)\n"
    "  Fatima   Score 2 – Very Urgent (Orange)\n"
    "  James    Score 3 – Urgent      (Yellow)\n"
    "  Marcus   Score 3 – Urgent      (Yellow)\n"
    "  Chen     Score 4 – Standard    (Green)\n"
    "  Oliver   Score 5 – Non-Urgent  (Blue)\n"
    "\n"
    "Critical patients flagged (score <= 2): 3\n"
    "  ** Maria    Score 1 – Immediate   (Red)\n"
    "  ** Priya    Score 1 – Immediate   (Red)\n"
    "  ** Fatima   Score 2 – Very Urgent (Orange)")
figure_caption(doc, "Figure 2: Programme Console Output for the Small Dataset (7 patients)")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – RESULTS AND TESTING
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "7. Results and Testing")

para(doc,
    "The programme was tested across three patient dataset sizes representing "
    "realistic ED operational scenarios: a seven-patient shift handover, a "
    "fifty-patient mid-shift snapshot, and a ten-thousand-record monthly audit "
    "export. All measurements were obtained on a standard Python 3.9 "
    "environment using time.perf_counter().")

# ── Small dataset ──────────────────────────────────────────────────────────
para(doc, "Small Dataset — 7 Patients (Shift Handover)", bold=True, space_before=6)

add_table(doc,
    ["Algorithm", "Time (ms)", "Comparisons", "Swaps", "Arithmetic", "Total Ops"],
    [
        ("Bubble Sort",          "0.0460", "21",  "9",  "28", "58"),
        ("Insertion Sort",       "0.0333", "14",  "9",  "20", "43"),
        ("Merge Sort",           "0.0406", "14",  "0",  "40", "54"),
        ("Quick Sort (3-way)",   "0.0423", "15", "13",  "37", "65"),
    ],
    "Table 2: Performance Results — Small Dataset (7 Patients)"
)

para(doc,
    "With only seven patients, all four algorithms complete in under one "
    "tenth of a millisecond and the timing differences are insignificant. "
    "Merge Sort records the fewest comparisons (14) and zero swaps, since its "
    "merge step writes elements directly into their final positions rather than "
    "exchanging them in place. Insertion Sort uses the fewest total operations "
    "(43), reflecting its efficiency advantage on small, bounded datasets. The "
    "system correctly identified three critical patients — Maria (Score 1), "
    "Priya (Score 1) and Fatima (Score 2) — demonstrating accurate triage "
    "flagging at the smallest scale.")

# ── Medium dataset ─────────────────────────────────────────────────────────
para(doc, "Medium Dataset — 50 Patients (Mid-Shift Snapshot)", bold=True, space_before=6)

add_table(doc,
    ["Algorithm", "Time (ms)", "Comparisons", "Swaps", "Arithmetic", "Total Ops"],
    [
        ("Bubble Sort",          "0.4800", "1,225", "418", "1,275", "2,918"),
        ("Insertion Sort",       "0.2973",   "466", "418",   "515", "1,399"),
        ("Merge Sort",           "0.2818",   "210",   "0",   "545",   "755"),
        ("Quick Sort (3-way)",   "0.2147",   "125",  "80",   "303",   "508"),
    ],
    "Table 3: Performance Results — Medium Dataset (50 Patients)"
)

para(doc,
    "At fifty patients the performance gap begins to emerge. Bubble Sort "
    "requires 2,918 total operations — nearly four times more than Merge "
    "Sort's 755. Notably, the three-way Quick Sort requires only 508 "
    "operations, outperforming Merge Sort, because it groups the many "
    "equal-score records together in a single partition pass rather than "
    "recursing through them pairwise. Of the fifty randomly generated "
    "patients, 24 (48%) were flagged as critical (score ≤ 2), reflecting a "
    "plausible high-acuity shift distribution (Sabah et al., 2023).")

# ── Large dataset ──────────────────────────────────────────────────────────
para(doc, "Large Dataset — 10,000 Patients (Monthly Audit Export)", bold=True, space_before=6)

add_table(doc,
    ["Algorithm", "Time (ms)", "Comparisons", "Swaps", "Arithmetic", "Total Ops"],
    [
        ("Bubble Sort",          "11,168.45", "49,995,000", "20,257,609", "50,005,000", "120,257,609"),
        ("Insertion Sort",        "4,978.86", "20,267,608", "20,257,609", "20,277,607",  "60,802,824"),
        ("Merge Sort",               "74.10",    "111,408",           "0",    "255,023",     "366,431"),
        ("Quick Sort (3-way)",        "41.96",     "23,951",      "13,956",     "55,879",      "93,786"),
    ],
    "Table 4: Performance Results — Large Dataset (10,000 Patients)"
)

para(doc,
    "The large-dataset results are decisive. Bubble Sort required "
    "120,257,609 total operations and took 11,168 milliseconds — over eleven "
    "seconds — making it completely unfit for clinical use at scale. Insertion "
    "Sort was twice as fast (4,979 ms) but still took nearly five seconds, "
    "which is unacceptable in a live ED queue. Merge Sort processed all ten "
    "thousand records in just 74 milliseconds with 366,431 operations — "
    "approximately 328 times faster than Bubble Sort.")

para(doc,
    "The most notable result is that the three-way Quick Sort was the fastest "
    "algorithm of all, completing in only 42 milliseconds with just 93,786 "
    "total operations — 76% fewer than Merge Sort. This is a direct "
    "consequence of the triage score distribution: with only five possible "
    "integer values spread across 10,000 records, each three-way partition "
    "eliminates on average 2,000 equal-score records from further recursion "
    "in a single pass, rather than recursing through them pairwise as Merge "
    "Sort does. This demonstrates that domain knowledge about the data "
    "distribution can significantly influence algorithm selection beyond "
    "theoretical Big-O analysis alone (Wibowo and Faisal, 2024). A total of "
    "3,993 patients (approximately 40%) were correctly flagged as critical "
    "across the large dataset.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "8. Conclusion")

para(doc,
    "This project successfully implemented and empirically evaluated four "
    "sorting algorithms — Bubble Sort, Insertion Sort, Merge Sort and "
    "three-way Quick Sort — within the context of an NHS emergency department "
    "patient triage management system. Testing across datasets of 7, 50 and "
    "10,000 patient records validated theoretical Big-O predictions and "
    "revealed domain-specific insights about algorithm behaviour on data with "
    "low cardinality.")

para(doc,
    "Merge Sort is recommended as the primary production algorithm because "
    "its guaranteed O(n log n) complexity in all cases — best, average and "
    "worst — means its performance is entirely predictable regardless of "
    "how patients happen to be distributed across the five triage categories "
    "on any given shift. It processed 10,000 records in 74 milliseconds, "
    "which satisfies the real-time responsiveness requirement of the system "
    "(Cormen et al., 2022). Its stability property — preserving the arrival "
    "order of patients with equal scores — is also clinically important, as "
    "it ensures fairness among patients of identical urgency.")

para(doc,
    "The three-way Quick Sort delivered the best measured performance on the "
    "large dataset (42 ms), owing to the highly repeated nature of triage "
    "scores. Its primary limitation — discovered during testing — is the "
    "risk of a RecursionError with a standard Lomuto pivot on deeply repeated "
    "values. While the randomised 3-way variant resolves this, its expected "
    "rather than guaranteed O(n log n) complexity makes it less suitable as "
    "the sole production algorithm for a safety-critical system. "
    "Bubble Sort and Insertion Sort are appropriate only for very small "
    "datasets (fewer than twenty patients) where implementation simplicity "
    "outweighs performance considerations (Shabbir et al., 2023).")

para(doc,
    "All three functional requirements were met: the sorted patient queue was "
    "correct across all dataset sizes; critical patients (score ≤ 2) were "
    "accurately flagged in every test; and full performance statistics "
    "including record count, elapsed time and operation breakdown were "
    "reported for every algorithm.")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 9 – FUTURE IMPROVEMENTS
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "9. Future Improvements")

para(doc,
    "Several enhancements would strengthen the system for production "
    "deployment. First, a real-time streaming mode using Python's heapq "
    "module could maintain a min-heap of patient scores, allowing new "
    "arrivals to be inserted in O(log n) time without re-sorting the entire "
    "queue. This would be more appropriate for continuous patient intake "
    "than batch sorting (Lafore, Broder and Canning, 2022).")

para(doc,
    "Second, replacing the flat text file with an SQLite database would "
    "enable persistent storage, structured querying and automatic indexing "
    "on the triage score column. SQL's built-in ORDER BY clause, backed by "
    "an index, could retrieve the sorted queue in sub-millisecond time for "
    "typical shift sizes. Third, a REST API built with Python's Flask "
    "framework would allow the triage sorter to be integrated directly into "
    "existing hospital information systems and accessed by bedside terminals "
    "or mobile nursing devices.")

para(doc,
    "Fourth, the system could be extended to incorporate a secondary sort "
    "key — for example, time of arrival — so that patients with identical "
    "triage scores are ordered by waiting time, ensuring that no patient is "
    "indefinitely delayed by a steady stream of equally urgent arrivals. "
    "Finally, a machine learning model trained on historical ED data could "
    "predict likely triage scores from presenting symptoms before formal "
    "nurse assessment, enabling early queue pre-positioning and reducing "
    "wait times for the most critical patients (Abuba et al., 2025).")


# ════════════════════════════════════════════════════════════════════════════
#  SECTION 10 – REFERENCES
# ════════════════════════════════════════════════════════════════════════════

heading(doc, "10. References")

refs = [
    ("Abuba, N.S., Baagyere, E.Y., Nakpih, C.I. and Wiredu, J.K., 2025. "
     "Optiflexsort: a hybrid sorting algorithm for efficient large-scale data "
     "processing. Journal of Advances in Mathematics and Computer Science, "
     "40(2), pp.67–81."),

    ("Cormen, T.H., Leiserson, C.E., Rivest, R.L. and Stein, C., 2022. "
     "Introduction to algorithms. 4th edn. Cambridge, MA: MIT Press."),

    ("Das, U., 2025. Python or Java in a data structures course? How about "
     "both? In 2025 ASEE Annual Conference & Exposition. Washington, DC: ASEE."),

    ("Fatima, P., 2023. Optimizing algorithm efficiency through advanced data "
     "structures in C++: a comparative analysis of performance, scalability "
     "and complexity. International Journal of Computations, Information and "
     "Manufacturing (IJCIM), 3(2), pp.66–72."),

    ("Lafore, R., Broder, A. and Canning, J., 2022. Data structures & "
     "algorithms in Python. Boston: Addison-Wesley Professional."),

    ("Mackway-Jones, K., Marsden, J. and Windle, J., 2014. Emergency triage: "
     "Manchester Triage Group. 3rd edn. Oxford: Wiley-Blackwell."),

    ("NHS England, 2024. A&E attendances and emergency admissions 2023–24 "
     "statistical commentary. Leeds: NHS England. Available at: "
     "https://www.england.nhs.uk/statistics/statistical-work-areas/ae-waiting-times-and-activity/ "
     "[Accessed: 13 June 2026]."),

    ("Reya, N.F., Ahmed, A., Zaman, T. and Islam, M.M., 2023. GreenPy: "
     "evaluating application-level energy efficiency in Python for green "
     "computing. Annals of Emerging Technologies in Computing (AETiC), "
     "7(3), pp.92–110."),

    ("Sabah, A.S., Abu-Naser, S.S., Helles, Y.E., Abdallatif, R.F., "
     "Samra, F.Y.A., Taha, A.H.A., Massa, N.M. and Hamouda, A.A., 2023. "
     "Comparative analysis of the performance of popular sorting algorithms "
     "on datasets of different sizes and characteristics. International "
     "Journal of Academic Engineering Research (IJAER), 7(6), pp.8–20."),

    ("Shabbir, A., Majeed, A., Iftikhar, M., Ali, R.H., Arshad, U., "
     "Shabbir, M.Z., Ijaz, A.Z., Ali, N. and Aftab, A., 2023. A review of "
     "algorithm complexities on different valued sorted and unsorted data. "
     "In 2023 International Conference on IT and Industrial Technologies "
     "(ICIT). Sialkot: IEEE, pp.1–6."),

    ("Wibowo, F.R. and Faisal, M., 2024. Comparative analysis of sorting "
     "algorithms: TimSort Python and classical sorting methods. JISA "
     "(Jurnal Informatika dan Sains), 7(1), pp.11–18."),
]

for ref in refs:
    reference_entry(doc, ref)


# ════════════════════════════════════════════════════════════════════════════
#  APPENDIX A – FULL PYTHON SOURCE CODE
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading(doc, "Appendix A: Full Python Source Code (triage_sorter.py)")

if os.path.exists("triage_sorter.py"):
    with open("triage_sorter.py", "r") as fh:
        code_block(doc, fh.read())

# ════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════

out = "CSAPSP_Individual_Report.docx"
doc.save(out)
print(f"Report saved → {out}")
